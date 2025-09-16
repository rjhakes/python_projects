import docx
import os, sys

def docxMerger(fileList):
    writer = docx.Document()
    os.system('cls')
    print(str(len(fileList)) + ' Files merged: ' + ', '.join(f[f.rfind(f'\\')+1:] for f in fileList) + '\n')
    for i in fileList:
        reader = docx.Document(i)
        for para in reader.paragraphs:
            newPara = writer.add_paragraph(para.text)
            if (para.style):
                newPara.style = para.style
        writer.add_page_break()
    return writer

def docxCreator(filePath, doc):
    doc.save(filePath)